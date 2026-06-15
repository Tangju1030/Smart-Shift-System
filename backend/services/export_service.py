"""导出服务 — 生成 docx 值班表和 CSV"""

import io
import csv
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from services.schedule_service import ScheduleService

WEEKDAYS = ['周一', '周二', '周三', '周四', '周五']
DUTY_SLOTS = ['第一节', '第二节', '第三节', '第四节']
WEEKDAY_CN = ['一', '二', '三', '四', '五']


class ExportService:

    def __init__(self, schedule_service: ScheduleService):
        self.svc = schedule_service

    def export_docx(self, week_no: int, week_dates: dict[str, str], slots: dict[str, dict[str, list[str]]]) -> bytes:
        """
        生成与参考文档格式一致的值班签到表 docx

        表格结构:
          行0(合并): 时段 | 星期一(日期) | 星期二(日期) | ...  (5列, 每列colspan=2)
          行1:       时段 | 姓名 | 签到 | 姓名 | 签到 | ...    (11列)
          行2-5:     第N节 | 姓名 / 姓名 | (空) | ...
        """
        doc = Document()

        # ── 页面设置 ──
        section = doc.sections[0]
        section.page_width = Cm(29.7)  # A4 横向
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        # ── 标题 ──
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f'第{week_no}周年委值班签到表')
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = '宋体'
        run.font.color.rgb = RGBColor(0, 0, 0)
        title.space_after = Pt(6)

        # ── 副标题（日期范围） ──
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dates_text = '  |  '.join(
            f'星期{WEEKDAY_CN[i]}（{week_dates.get(wd, "")}）'
            for i, wd in enumerate(WEEKDAYS)
            if week_dates.get(wd, '') != '假'
        )
        run2 = subtitle.add_run(dates_text)
        run2.font.size = Pt(10)
        run2.font.name = '宋体'
        run2.font.color.rgb = RGBColor(100, 100, 100)
        subtitle.space_after = Pt(12)

        # ── 构建表格 ──
        # 列结构: col0=时段, 然后每工作日2列(姓名+签到)
        num_cols = 1 + len(WEEKDAYS) * 2  # 11列
        num_rows = len(DUTY_SLOTS) + 2    # 4数据行 + 2表头行 = 6行

        table = doc.add_table(rows=num_rows, cols=num_cols, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # ── 设置列宽 ──
        col_widths = [Cm(2.2)]  # 时段列
        for _ in WEEKDAYS:
            col_widths.append(Cm(2.0))   # 姓名列
            col_widths.append(Cm(1.5))   # 签到列
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

        # ── 行0: 表头（合并单元格） ──
        # 列0: "时段" 垂直合并行0和行1
        cell_00 = table.cell(0, 0)
        cell_00.text = ''
        self._set_cell_text(cell_00, '时段', bold=True, size=10, center=True)
        self._merge_v(cell_00, table.cell(1, 0))

        # 每个工作日占2列，行0合并
        for wi, wd in enumerate(WEEKDAYS):
            date = week_dates.get(wd, '')
            label = f'星期{WEEKDAY_CN[wi]}（{date}）' if date != '假' else f'星期{WEEKDAY_CN[wi]}（放假）'
            col_name = 1 + wi * 2
            cell_a = table.cell(0, col_name)
            cell_b = table.cell(0, col_name + 1)
            self._set_cell_text(cell_a, label, bold=True, size=10, center=True)
            self._merge_h(cell_a, cell_b)

        # ── 行1: 子表头（姓名 / 签到） ──
        self._set_cell_text(table.cell(1, 0), '时段', bold=True, size=9, center=True)
        for wi in range(len(WEEKDAYS)):
            col_name = 1 + wi * 2
            self._set_cell_text(table.cell(1, col_name), '姓名', bold=True, size=9, center=True)
            self._set_cell_text(table.cell(1, col_name + 1), '签到', bold=True, size=9, center=True)

        # ── 行2-5: 数据行 ──
        for si, slot in enumerate(DUTY_SLOTS):
            row_idx = si + 2
            row = table.rows[row_idx]
            # 时段名
            self._set_cell_text(row.cells[0], slot, bold=True, size=10, center=True)
            # 数据
            for wi, wd in enumerate(WEEKDAYS):
                col_name = 1 + wi * 2
                col_sign = col_name + 1
                if week_dates.get(wd, '') == '假':
                    self._set_cell_text(row.cells[col_name], '放假', size=9, center=True)
                    self._set_cell_text(row.cells[col_sign], '', size=9, center=True)
                else:
                    names = slots.get(wd, {}).get(slot, [])
                    self._set_cell_text(row.cells[col_name], ' / '.join(names) if names else '', size=9)
                    self._set_cell_text(row.cells[col_sign], '', size=9)

        # ── 设置所有单元格垂直居中 ──
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
                tcPr.append(vAlign)

        # ── 统计信息 ──
        doc.add_paragraph()  # 空行
        stats_para = doc.add_paragraph()
        stats_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 收集统计
        all_names = set()
        total_assignments = 0
        for wd in WEEKDAYS:
            for slot in DUTY_SLOTS:
                names = slots.get(wd, {}).get(slot, [])
                all_names.update(names)
                total_assignments += len(names)
        stats_text = (
            f'本周值班总人次：{total_assignments}  |  '
            f'参与人数：{len(all_names)}  |  '
            f'时段：第一节 8:30~9:50 / 第二节 10:25~11:50 / 第三节 14:30~15:55 / 第四节 16:30~17:40'
        )
        run3 = stats_para.add_run(stats_text)
        run3.font.size = Pt(9)
        run3.font.name = '宋体'
        run3.font.color.rgb = RGBColor(120, 120, 120)

        # ── 输出 ──
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    def export_csv(self, week_no: int, week_dates: dict, slots: dict) -> str:
        """导出值班表视图 CSV"""
        output = io.StringIO()
        output.write('﻿')  # BOM for Excel
        writer = csv.writer(output)

        header = ['时段']
        for wd in WEEKDAYS:
            date = week_dates.get(wd, '')
            header.append(f'{wd}({date})')
        writer.writerow(header)

        for slot in DUTY_SLOTS:
            row = [slot]
            for wd in WEEKDAYS:
                names = slots.get(wd, {}).get(slot, [])
                row.append(' / '.join(names) if names else '')
            writer.writerow(row)

        return output.getvalue()

    # ── 辅助方法 ──

    def _set_cell_text(self, cell, text: str, bold: bool = False, size: int = 10, center: bool = False):
        """设置单元格文字"""
        # 清除已有段落
        for p in cell.paragraphs:
            p.clear()
        p = cell.paragraphs[0]
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(text))
        run.font.size = Pt(size)
        run.font.name = '宋体'
        run.bold = bold

    def _merge_h(self, cell_a, cell_b):
        """水平合并 cell_a 和 cell_b"""
        tc_a = cell_a._tc
        tc_b = cell_b._tc
        tcPr_a = tc_a.get_or_add_tcPr()
        tcPr_b = tc_b.get_or_add_tcPr()
        # gridSpan on cell_a
        grid_span = parse_xml(f'<w:gridSpan {nsdecls("w")} w:val="2"/>')
        tcPr_a.append(grid_span)
        # merge on both
        merge_a = parse_xml(f'<w:hMerge {nsdecls("w")} w:val="restart"/>')
        merge_b = parse_xml(f'<w:hMerge {nsdecls("w")} w:val="continue"/>')
        tcPr_a.append(merge_a)
        tcPr_b.append(merge_b)

    def _merge_v(self, cell_top, cell_bottom):
        """垂直合并 cell_top 和 cell_bottom"""
        tcPr_top = cell_top._tc.get_or_add_tcPr()
        tcPr_bottom = cell_bottom._tc.get_or_add_tcPr()
        merge_top = parse_xml(f'<w:vMerge {nsdecls("w")} w:val="restart"/>')
        merge_bottom = parse_xml(f'<w:vMerge {nsdecls("w")} w:val="continue"/>')
        tcPr_top.append(merge_top)
        tcPr_bottom.append(merge_bottom)
